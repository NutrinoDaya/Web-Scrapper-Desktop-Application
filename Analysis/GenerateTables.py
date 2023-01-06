from docx import Document
from docx.shared import Cm, Pt,Inches
import pathlib
import os
# Import date class from datetime module
from datetime import datetime as dt


desktop = pathlib.Path.home() / 'Desktop'
desktop = str(desktop)
desktop+= "\\Generated Data\\"
GNDataFolder = desktop

def CreateTable(SortedNames,Sorted_Citations,FileNum,keycode):


    doc_Files = ["Data_Dep_1.docx", "Data_Dep_2.docx","Data_Dep_3.docx","Data_Dep_4.docx",
                 "Data_Dep_5.docx","Data_Dep_6.docx","Data_Dep_7.docx",
                 "Data_Dep_8.docx","Data_Dep_9.docx"]

    Departments = ["Information Technology Department", "Electrical Engineering Department",
                   "Water and Environmental Department",
                   "Mechanical Engineering Department", "Chemical Engineering Department",
                   "Nutrition and Food Processing Department",
                   "Vocational Education Department", "Department of Basic Human Sciences",
                   "Department of Basic Scientific Sciences"]

    Pictures = ["DataAnalysis1.png", "DataAnalysis2.png","DataAnalysis3.png","DataAnalysis4.png","DataAnalysis5.png",
                "DataAnalysis6.png","DataAnalysis7.png","DataAnalysis8.png","DataAnalysis9.png"]

    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
    from docx import Document
    from docx.shared import Cm, Pt, Inches

    if keycode == 0 :
        SortedNames.reverse()
        Sorted_Citations.reverse()

        word_document = Document()
        word_document.add_heading(Departments[FileNum],1)
        # Number of Cells and number of columns
        table = word_document.add_table(2,3)  # we add rows iteratively

        table.style = 'Table Grid'

        # Changing Columns Width
        col = table.columns[0]
        col.width = Inches(0.5)
        # For Name  Cell
        cell = table.rows[0].cells[0]
        cell.text = "Citations"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True
        # For Citations Cell
        cell = table.rows[0].cells[1]

        cell.text = "Name"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True
        cell = table.rows[0].cells[2]
        cell.text = "Number"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True


        # table.add_row()

        for i in range(1,len(Sorted_Citations)+1):
            table.add_row()
            currRow = table.rows[i]
            currRow.cells[0].text = str(Sorted_Citations[i-1])

        for i in range(1, len(SortedNames)+1):
            currRow = table.rows[i]
            currRow.cells[1].text = SortedNames[i-1]

        for i in range(1, len(SortedNames)+1):
            currRow = table.rows[i]
            currRow.cells[2].text = str(i)
        # Delete The Last Empty Row
        row = table.rows[-1]
        remove_row(table, row)
        word_document.add_paragraph().add_run(
            ' ')
        word_document.add_paragraph().add_run(
            ' ')
        # print("FileNum = " , FileNum)
        # print("Curr Picture = " ,Pictures[FileNum])

        word_document.add_picture(GNDataFolder+Pictures[FileNum])
        word_document.add_paragraph().add_run(
            '\n')
        # Returns the current local date
        now = dt.now() # current date and time
        date_time = 'This Report Was Created On  :'
        date_time += now.strftime("%m/%d/%Y")
        word_document.add_paragraph(
        date_time,style='Intense Quote'
        )
        word_document.add_page_break()

        # word_document.add_page_break()
        # print("FileNum = " , FileNum)m

        word_document.save(GNDataFolder+Departments[FileNum]+".docx")
        # print("Word File Saved")
    else :
        # print("Sorted_Citations = " , Sorted_Citations)
        # print("1 Key Code Called ")
        # SortedNames.reverse()
        # Sorted_Citations.reverse()

        word_document = Document()

        table = word_document.add_table(2, 3)  # we add rows iteratively

        table.style = 'Table Grid'

        # Changing Columns Width
        col = table.columns[0]
        col.width = Inches(0.5)
        # For Name  Cell
        cell = table.rows[0].cells[0]
        cell.text = "Citations"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True
        # For Citations Cell
        cell = table.rows[0].cells[1]
        cell.text = "Department"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True
        cell = table.rows[0].cells[2]
        cell.text = "Number"
        cell.paragraphs[0].runs[0].font.size = Pt(14)
        cell.paragraphs[0].runs[0].font.bold = True

        # table.add_row()

        for i in range(1, len(Sorted_Citations) + 1):
            table.add_row()
            currRow = table.rows[i]
            currRow.cells[0].text = str(Sorted_Citations[i - 1])

        for i in range(1, len(SortedNames) + 1):
            currRow = table.rows[i]
            currRow.cells[1].text = SortedNames[i - 1]
        # Delete The Last Empty Row
        for i in range(1, len(SortedNames)+1):
            currRow = table.rows[i]
            currRow.cells[2].text = str(i)
        row = table.rows[-1]
        remove_row(table, row)
        word_document.add_paragraph().add_run(
            ' ')
        word_document.add_paragraph().add_run(
            ' ')

        word_document.add_picture(GNDataFolder + "LastPlot.png")
        word_document.add_paragraph().add_run(
            '\n')
        now = dt.now() # current date and time
        date_time = 'This Report Was Created On  :'
        date_time += now.strftime("%m/%d/%Y")
        word_document.add_paragraph(
        date_time,style='Intense Quote'
        )
        word_document.add_page_break()

        # word_document.add_page_break()
        word_document.save(GNDataFolder + "All Departments Report.docx")
        # print("Word File Saved")